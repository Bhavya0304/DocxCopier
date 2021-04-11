import docx
import sys

try:
    user_arg = sys.argv[1]
except Exception as e:
    print("Please Enter the file name to be edited as argument... (Only Docx Format)")
    sys.exit()



print("Enter all details to be replaced...")
usr1 = input("Exact Name written in provided file: ")
usr2 = input("Exact Enroll written in provided file: ")
usr3 = input("Exact Division written in provided file: ")

main_list = []


#Get all style values
def getvalue(pH):    
    bold = pH.bold
    italic = pH.italic
    fontsize = pH.font.size
    fontcolor = pH.font.color.rgb
    return [bold,italic,fontsize,fontcolor]
#Set all style values
def addvalue(run,list1):
    run.bold = list1[0]
    run.italic = list1[1]
    run.font.size = list1[2]
    run.font.color.rgb = list1[3]
    return run



while(True):
    usr_list = []
    for i in range(0,3):
        if(i == 0):
            inp = input("Enter New Name: ")
        elif(i == 1):
            inp = input("Enter New Enroll: ")
        elif(i == 2):
            inp = input("Enter New Division: ")
        usr_list.append(inp)
    main_list.append(usr_list)
    inp = input("Enter to Add More (press e and enter to exit the adding process!): ")
    if(inp == "e"):
        break
    else:
        continue    
    

subject = input("State Subject Name: ")
pracassign = input("State Practicals or Assignemnts in format(Practicals(1-7) or Assignments(1-2)): ")    
print("Creating Lists of Profiles...")
print(main_list)               




for user in main_list:
    #Managing Header Sections
    print("\n\n Changing Header portion..(if required)...")
    doc = docx.Document(user_arg)
    section = doc.sections[0]
    header = section.header
    for para in header.paragraphs:
        for run in para.runs:
            if(run.text.find(usr2) != -1):
                run_vals = getvalue(run)
                run.text = run.text.replace(usr2,user[1])
                run = addvalue(run,run_vals)
            else:
                pass    
       
    #Managing Data Entires
    print("\n\n Changing file Data....")
    for para in doc.paragraphs:
        for run in para.runs:
            if(run.text.find(usr1) != -1):
                run_vals = getvalue(run)
                run.text = run.text.replace(usr1,user[0])
                run = addvalue(run,run_vals)
            if(run.text.find(usr2) != -1):
                run_vals = getvalue(run)
                run.text = run.text.replace(usr2,user[1])
                run = addvalue(run,run_vals)
            if(run.text.find(usr3) != -1):
                run_vals = getvalue(run)
                run.text = run.text.replace(usr3,user[2])
                run = addvalue(run,run_vals)
    print("\n\n Changing data in tables...(if required)...")            
    

    print("Creating files...")
    name = f"{user[1]}_{subject}_{pracassign}.docx"      
    doc.save(name)


